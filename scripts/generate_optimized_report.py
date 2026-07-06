#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Generate optimized Word report for the highway lane-change DQN vs PPO study.
Creates a properly formatted .docx with A4 page, page numbers, proper fonts,
heading styles, formatted tables, and embedded figures.

Key formatting:
- A4 paper, margins: top/bottom 2.54cm, left/right 2.8cm
- Chinese font: 宋体 (SimSun), English font: Times New Roman
- Body text: 小四 (12pt), 1.5x line spacing, first-line indent 2 chars
- Title: 黑体 (SimHei), 三号 (16pt), bold, centered
- Heading 1: 黑体, 小三 (15pt), bold
- Heading 2: 黑体, 小四 (12pt), bold
- Heading 3: 黑体, 小四 (12pt), bold
- Page numbers: footer, centered
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

# Paths
PROJECT_ROOT = r"D:\highway-lane-change"
FIGURES_DIR = os.path.join(PROJECT_ROOT, "figures")
OUTPUT_PATH = os.path.join(PROJECT_ROOT, "final_submit", "基于DQN与PPO的高速公路自主变道决策研究_优化版.docx")

# ============================================================
# Helper functions
# ============================================================

def set_cell_font(cell, text, font_name_cn='宋体', font_name_en='Times New Roman',
                  size=Pt(9), bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    """Set cell text with proper Chinese/English fonts."""
    # Clear existing paragraphs
    for p in cell.paragraphs:
        for r in p.runs:
            r.clear()
    p = cell.paragraphs[0]
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(str(text))
    run.font.size = size
    run.bold = bold
    run.font.name = font_name_en
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name_cn)
    return run


def set_run_font(run, font_name_cn='宋体', font_name_en='Times New Roman', size=Pt(12)):
    """Set font for a run with both Chinese and English font names."""
    run.font.name = font_name_en
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name_cn)
    run.font.size = size


def add_paragraph_with_font(doc, text, font_name_cn='宋体', font_name_en='Times New Roman',
                            size=Pt(12), bold=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                            first_line_indent=None, space_before=Pt(0), space_after=Pt(0),
                            line_spacing=1.5):
    """Add a paragraph with proper Chinese/English font settings."""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = space_before
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = line_spacing
    if first_line_indent:
        p.paragraph_format.first_line_indent = first_line_indent
    run = p.add_run(text)
    run.font.name = font_name_en
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name_cn)
    run.font.size = size
    run.bold = bold
    return p


def add_heading_styled(doc, text, level=1):
    """Add a heading with proper Chinese/English font settings."""
    h = doc.add_heading(text, level=level)
    # Set fonts for all runs in heading
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        if level == 0:
            run.font.size = Pt(16)
        elif level == 1:
            run.font.size = Pt(15)
        elif level == 2:
            run.font.size = Pt(12)
        elif level == 3:
            run.font.size = Pt(12)
    return h


def add_body_text(doc, text):
    """Add body text paragraph with standard formatting: 宋体/TNR, 小四, 1.5x spacing, 2-char indent."""
    return add_paragraph_with_font(
        doc, text,
        font_name_cn='宋体', font_name_en='Times New Roman',
        size=Pt(12), bold=False,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Cm(0.74),  # approximately 2 Chinese characters
        space_before=Pt(0), space_after=Pt(0),
        line_spacing=1.5
    )


def add_centered_text(doc, text, size=Pt(12), bold=False, font_cn='宋体', font_en='Times New Roman'):
    """Add centered text paragraph."""
    return add_paragraph_with_font(
        doc, text,
        font_name_cn=font_cn, font_name_en=font_en,
        size=size, bold=bold,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=None,
        space_before=Pt(3), space_after=Pt(3),
        line_spacing=1.5
    )


def add_figure_with_caption(doc, image_path, caption_text, width_cm=14.0):
    """Add a centered figure with caption below it."""
    if os.path.exists(image_path):
        # Add the image
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(6)
        p_img.paragraph_format.space_after = Pt(3)
        run = p_img.add_run()
        run.add_picture(image_path, width=Cm(width_cm))
        # Add caption
        add_paragraph_with_font(
            doc, caption_text,
            font_name_cn='宋体', font_name_en='Times New Roman',
            size=Pt(9), bold=False,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent=None,
            space_before=Pt(0), space_after=Pt(6),
            line_spacing=1.2
        )
    else:
        add_body_text(doc, f"[图片缺失: {os.path.basename(image_path)}]")


def create_styled_table(doc, headers, rows, col_widths=None, header_size=Pt(9), body_size=Pt(9)):
    """Create a properly formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Format header
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_font(cell, header, size=header_size, bold=True,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
        # Set header background
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3"/>')
        cell._element.get_or_add_tcPr().append(shading)

    # Format body rows
    for r, row_data in enumerate(rows):
        for c, cell_text in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            align = WD_ALIGN_PARAGRAPH.CENTER if c > 0 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_font(cell, cell_text, size=body_size, bold=False, alignment=align)

    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Cm(width)

    # Add spacing before and after table
    return table


def add_page_number(doc):
    """Add page numbers to footer (centered)."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

        # Add page number field
        run = p.add_run()
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._element.append(fldChar1)

        run2 = p.add_run()
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._element.append(instrText)

        run3 = p.add_run()
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._element.append(fldChar2)

        set_run_font(run, '宋体', 'Times New Roman', Pt(9))
        set_run_font(run2, '宋体', 'Times New Roman', Pt(9))
        set_run_font(run3, '宋体', 'Times New Roman', Pt(9))


def set_default_style(doc):
    """Set the document default style."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


# ============================================================
# Main document creation
# ============================================================

def create_report():
    doc = Document()

    # --- Page setup ---
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)

    set_default_style(doc)

    # --- Configure heading styles ---
    for i, (size, font_cn) in enumerate([(Pt(16), '黑体'), (Pt(15), '黑体'), (Pt(12), '黑体'), (Pt(12), '黑体')]):
        style_name = f'Heading {i}' if i > 0 else 'Title'
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        style.font.name = 'Times New Roman'
        style.font.size = size
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), font_cn)
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(6)

    # ============================
    # TITLE
    # ============================
    add_paragraph_with_font(
        doc, '基于 DQN 与 PPO 的高速公路自主变道决策研究',
        font_name_cn='黑体', font_name_en='Times New Roman',
        size=Pt(16), bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=Pt(30), space_after=Pt(20),
        line_spacing=1.5
    )

    # ============================
    # 中文摘要
    # ============================
    add_heading_styled(doc, '中文摘要', level=1)

    abstract_cn = (
        '高速公路自主变道是自动驾驶汽车安全高效行驶的关键技术之一，涉及安全性、通行效率、乘坐舒适性'
        '和泛化能力等多个相互冲突的优化目标，构成一类典型的复杂工程问题。本研究基于 highway-env 仿真平台'
        '构建高速公路多车交互环境，将变道决策建模为强化学习（Reinforcement Learning, RL）中的序贯决策问题，'
        '以深度 Q 网络（Deep Q-Network, DQN）和近端策略优化（Proximal Policy Optimization, PPO）两种代表性'
        '深度强化学习算法为核心开展对比研究。通过设计包含碰撞惩罚、速度激励、变道代价、急变速代价和靠右行驶'
        '奖励五个分量的多目标奖励函数，系统比较了 baseline、comfort、aggressive、balanced、conservative '
        '五种权重配置。实验采用 OccupancyGrid 状态观测与 MlpPolicy 策略网络，共完成 2 种算法 × 5 种奖励配置'
        ' × 5 种随机种子合计 50 组训练，每组算法-奖励组合选出最优模型，在分布内场景及车道关闭、前车急刹、'
        '高密度车流三种分布外（Out-of-Distribution, OOD）场景下进行泛化评估。结果表明：PPO 算法在收敛速度、'
        '训练稳定性和最终性能方面整体优于 DQN；PPO aggressive 策略在四个场景下平均成功率达到 97.5%、'
        '碰撞率仅 2.5%、平均回报 163.40，为综合最优策略；DQN 存在严重的频繁变道和策略不稳定问题，碰撞率'
        '高达 94% 以上。研究进一步揭示：高回报不一定等于安全策略，必须结合成功率、碰撞率和变道次数等'
        '多维指标进行综合评价。'
    )
    add_body_text(doc, abstract_cn)

    # Keywords
    kw_cn = '关键词：高速公路；自主变道；深度强化学习；DQN；PPO；多目标奖励塑形；分布外泛化'
    add_paragraph_with_font(
        doc, kw_cn,
        font_name_cn='宋体', font_name_en='Times New Roman',
        size=Pt(12), bold=True,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Cm(0.74),
        space_before=Pt(6), space_after=Pt(6),
        line_spacing=1.5
    )

    # ============================
    # English Abstract
    # ============================
    add_heading_styled(doc, 'English Abstract', level=1)

    abstract_en = (
        'Highway autonomous lane-changing is a critical technology for safe and efficient self-driving. '
        'This paper models lane-changing decision-making as a sequential decision problem in reinforcement '
        'learning using the highway-env simulation platform. Two representative deep reinforcement learning '
        'algorithms, Deep Q-Network (DQN) and Proximal Policy Optimization (PPO), are systematically compared '
        'under five multi-objective reward configurations: baseline, comfort, aggressive, balanced, and '
        'conservative. A total of 50 training runs (2 algorithms × 5 rewards × 5 seeds) are conducted '
        'using OccupancyGrid observations and MlpPolicy networks. The best model from each algorithm-reward '
        'combination is evaluated under four scenarios, including three out-of-distribution (OOD) settings: '
        'lane closure, sudden braking, and high-density traffic. Experimental results demonstrate that PPO '
        'consistently outperforms DQN in convergence speed, training stability, and final performance. '
        'The PPO aggressive strategy achieves the best overall performance with an average success rate of '
        '97.5%, collision rate of 2.5%, and mean reward of 163.40 across all scenarios. DQN exhibits severe '
        'policy instability and frequent unnecessary lane changes, with collision rates exceeding 94%. '
        'A key finding is that high mean reward alone does not guarantee a safe policy: comprehensive '
        'evaluation using success rate, collision rate, and lane-change frequency is essential. This study '
        'also discusses the social, cultural, regulatory, and ethical dimensions of autonomous driving '
        'deployment across different national contexts.'
    )
    add_body_text(doc, abstract_en)

    kw_en = 'Keywords: Highway; Autonomous Lane Change; Deep Reinforcement Learning; DQN; PPO; Multi-Objective Reward Shaping; Out-of-Distribution Generalization'
    add_paragraph_with_font(
        doc, kw_en,
        font_name_cn='宋体', font_name_en='Times New Roman',
        size=Pt(12), bold=True,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Cm(0.74),
        space_before=Pt(6), space_after=Pt(6),
        line_spacing=1.5
    )

    # ================================================================
    # 1 引言
    # ================================================================
    add_heading_styled(doc, '1 引言', level=1)

    intro_paragraphs = [
        '高速公路场景下的自主变道决策是自动驾驶领域最具挑战性的问题之一。与城市道路不同，高速公路车速高、'
        '车流密集、反应时间短，变道决策需要在极短时间内综合考虑自车状态、周围车辆位置与速度、道路几何结构'
        '以及交通规则等多维度信息。一次错误的变道决策可能导致严重交通事故，危及乘客及其他道路使用者的生命安全。',

        '该问题之所以成为复杂工程问题，在于其具备以下特征。第一，优化目标多元且相互冲突：安全性要求避免碰撞、'
        '效率要求保持高速度、舒适性要求减少急变速和频繁变道。第二，环境具有高度不确定性和动态性：周围车辆的'
        '行为不可预测，交通密度持续变化，还可能面临车道关闭、前车急刹等突发状况。第三，决策具有序贯性：当前'
        '变道决策不仅影响当前时刻的安全和效率，还会影响未来的状态和可选动作。因此，高速公路自主变道并非简单'
        '的分类或回归问题，而是在连续交互中不断优化长期累积回报的序贯决策问题。',

        '强化学习作为机器学习的重要范式，通过智能体与环境的持续交互和试错学习，天然适用于此类序贯决策问题。'
        '深度强化学习将深度神经网络的表征能力与强化学习的决策优化能力相结合，已在游戏博弈、机器人控制和'
        '自动驾驶等领域展现出强大潜力。本研究选取两种代表性深度强化学习算法——基于值函数的深度 Q 网络（DQN）'
        '和基于策略梯度的近端策略优化（PPO）——在统一仿真环境下进行系统比较，旨在回答以下问题：在高速公路'
        '变道决策任务中，哪类算法更为适用？不同奖励函数设计如何影响驾驶策略的安全性和效率？模型在未见过的'
        '交通场景下泛化能力如何？',

        '本研究属于《模式识别与机器学习》课程中的机器学习方法综合应用。通过完整的"问题建模—算法选择—'
        '仿真实验—结果分析—结论总结"流程，考核面向复杂工程问题选择或设计合适的模式识别与机器学习方法、'
        '正确处理数据、实现算法、开展验证并得到结果、分析并形成结论的能力，突出复杂工程问题求解能力的培养。'
    ]
    for para in intro_paragraphs:
        add_body_text(doc, para)

    # ================================================================
    # 2 复杂工程问题分析
    # ================================================================
    add_heading_styled(doc, '2 复杂工程问题分析', level=1)

    add_body_text(doc, '本研究符合复杂工程问题的七个核心特征，具体分析如下。')

    features = [
        ('特征一：问题需要深入分析才能确定解决方法。',
         '高速公路自主变道并非一眼就能确定解决方法的问题。需要分析多个优化目标（安全、效率、舒适性）、'
         '多种约束条件（交通规则、车辆动力学、传感器范围）和多个评价指标（成功率、碰撞率、速度、变道次数），'
         '在此基础上选择或设计合适的算法框架。'),
        ('特征二：涉及多方面技术因素。',
         '本问题涉及感知状态表征（OccupancyGrid 网格观测）、决策控制策略（离散元动作空间）、奖励函数设计'
         '（五分量加权）、算法训练稳定性（多种子重复实验）、安全性评估（碰撞率）、舒适性（变道次数、'
         '急变速次数）和交通效率（平均速度）等多个维度的技术因素。'),
        ('特征三：需要建立抽象模型。',
         '本研究将高速公路变道决策抽象为马尔可夫决策过程（Markov Decision Process, MDP），定义状态空间 S'
         '（OccupancyGrid 观测）、动作空间 A（五种离散元动作）、状态转移概率 P（仿真环境动力学）和奖励函数 R'
         '（多目标加权奖励），将工程问题转化为强化学习中的策略优化问题。'),
        ('特征四：不是单一常规方法即可解决。',
         '本研究比较了 DQN（基于值函数、off-policy）和 PPO（基于策略梯度、on-policy）两类不同范式的深度'
         '强化学习算法，同时比较了五种 reward 配置（baseline、comfort、aggressive、balanced、conservative），'
         '体现了多方法、多配置的系统对比研究思路。'),
        ('特征五：涉及不确定性、动态性和非线性。',
         '高速公路交通环境具有高度动态性和不确定性：周围车辆的行为不可预测、交通密度持续变化、车道关闭和'
         '突然刹车等突发事件随时可能发生。车辆动力学和碰撞物理也具有非线性特征。模型需要在高度不确定的环境'
         '中做出实时决策。'),
        ('特征六：涉及多方利益。',
         '自动驾驶变道决策涉及多方利益相关者：自车乘客（追求安全和舒适）、道路其他车辆（追求不被干扰和'
         '碰撞风险最小化）、交通管理者（追求整体通行效率最大化）、系统开发者（追求算法可靠性和合规性）。'
         '不同利益方之间存在天然的张力：例如提高自车速度可能增加其他车辆的安全风险，增加变道次数可能提高'
         '通行效率但降低乘坐舒适性。'),
        ('特征七：具有系统性。',
         '本研究包含多个相互关联的子系统：仿真环境构建（highway-env 配置与包装）、奖励函数设计'
         '（utils/reward.py）、强化学习训练流水线（train_dqn.py、train_ppo.py）、模型评估（eval.py 与 OOD '
         '场景测试）、最优模型选择（select_best.py）、结果可视化（scripts/plot_*.py）和日志追踪系统'
         '（utils/logging.py）。各模块协同工作，构成完整的实验研究系统。'),
    ]
    for title, body in features:
        add_paragraph_with_font(
            doc, title, font_name_cn='黑体', font_name_en='Times New Roman',
            size=Pt(12), bold=True,
            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            first_line_indent=Cm(0.74),
            space_before=Pt(3), space_after=Pt(0),
            line_spacing=1.5
        )
        add_body_text(doc, body)

    add_body_text(doc,
        '综上所述，本研究并非简单地"运行一个强化学习实验"，而是系统性地解决了一个具有多目标冲突、'
        '环境不确定性、多方利益权衡和工程部署约束的复杂工程问题。'
    )

    # ================================================================
    # 3 算法介绍
    # ================================================================
    add_heading_styled(doc, '3 算法介绍', level=1)

    # 3.1
    add_heading_styled(doc, '3.1 强化学习基础', level=2)

    add_body_text(doc,
        '强化学习研究智能体（Agent）如何在环境（Environment）中通过试错学习最优策略。在每个时间步 t，'
        '智能体观测环境状态 s_t，选择动作 a_t，环境返回即时奖励 r_t 并转移到新状态 s_{t+1}。智能体的'
        '目标是学习一个策略 π(a|s)，使得从当前状态开始的累积折扣回报 G_t = Σ_{k=0}^{∞} γ^k r_{t+k} '
        '最大化，其中 γ ∈ [0,1] 为折扣因子。'
    )
    add_body_text(doc,
        '在本项目中，状态 s 为 OccupancyGrid 观测（16×16 网格 + 速度特征），动作 a 为五种离散化的'
        '高层元动作（左变道、右变道、加速、减速、保持），奖励 r 为五分量加权和，γ = 0.8。'
    )

    # 3.2
    add_heading_styled(doc, '3.2 深度 Q 网络（DQN）', level=2)

    add_body_text(doc,
        'DQN 由 Mnih 等人于 2015 年提出 [1]，开创性地将深度神经网络与 Q-learning 相结合。其核心思想'
        '是使用深度神经网络 Q(s, a; θ) 近似最优动作价值函数 Q*(s, a)，并采用 ε-greedy 策略选择动作：'
        '以概率 1-ε 选择 Q 值最大的动作，以概率 ε 随机探索。'
    )
    add_body_text(doc,
        'DQN 引入两项关键技术解决训练不稳定问题：（1）经验回放（Experience Replay），将智能体的交互'
        '经验 (s, a, r, s\') 存入回放缓冲区，训练时随机采样小批量数据，打破样本间的时序相关性；'
        '（2）目标网络（Target Network），维护一个参数为 θ^- 的独立目标网络，周期性从主网络复制参数，'
        '使 Q-learning 目标 y = r + γ max_{a\'} Q(s\', a\'; θ^-) 在短时间内保持稳定。'
    )
    add_body_text(doc,
        '本研究中 DQN 的超参数配置为：学习率 5×10^{-4}，缓冲区大小 50000，批大小 64，折扣因子 '
        'γ = 0.8，目标网络软更新系数 τ = 0.005，初始探索率 ε = 1.0，最终探索率 ε = 0.05。'
    )

    # 3.3
    add_heading_styled(doc, '3.3 近端策略优化（PPO）', level=2)

    add_body_text(doc,
        'PPO 由 Schulman 等人于 2017 年提出 [2]，是目前应用最广泛的策略梯度算法之一。与 DQN 学习 '
        'Q 值函数不同，PPO 直接优化策略网络 π_θ(a|s)，采用 Actor-Critic 架构：Actor 网络输出动作的'
        '概率分布，Critic 网络估计状态价值函数 V(s)，用于计算优势函数 A(s, a)。'
    )
    add_body_text(doc,
        'PPO 的核心创新是使用裁剪代理目标函数（Clipped Surrogate Objective）限制新旧策略之间的差异：'
    )
    # PPO formula
    add_paragraph_with_font(
        doc, 'L^{CLIP}(θ) = E_t[min(r_t(θ)A_t, clip(r_t(θ), 1-ε, 1+ε)A_t)]',
        font_name_cn='宋体', font_name_en='Times New Roman',
        size=Pt(12), bold=False,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=Pt(3), space_after=Pt(3),
        line_spacing=1.5
    )
    add_body_text(doc,
        '其中 r_t(θ) = π_θ(a_t|s_t) / π_{θ_old}(a_t|s_t) 为新旧策略的概率比值，ε = 0.2 为裁剪范围。'
        '当优势函数为正时，限制策略更新幅度不超过 1+ε；当优势函数为负时，限制策略更新幅度不低于 1-ε。'
        '这一机制有效防止了策略更新过大导致的训练崩溃。'
    )
    add_body_text(doc,
        'PPO 属于在线（on-policy）算法，每轮使用当前策略采集一批新样本后进行多轮小批量优化。'
    )
    add_body_text(doc,
        '本研究中 PPO 的超参数配置为：学习率 3×10^{-4}，步数 512，批大小 64，优化轮数 10，折扣因子 '
        'γ = 0.8，GAE λ = 0.95，clip 范围 0.2，熵系数 0.01。'
    )

    # 3.4
    add_heading_styled(doc, '3.4 两种算法对比的意义', level=2)

    add_body_text(doc,
        'DQN（off-policy）和 PPO（on-policy）代表了深度强化学习的两种主流范式 [4][5]。DQN 理论上可以'
        '利用历史数据进行学习，样本效率较高，但在高维状态空间下 Q 值函数估计可能不准确，且离散动作空间'
        '中的 max 操作可能引入过高估计偏差。PPO 通过限制更新幅度保证训练稳定性，在连续控制和高维观测'
        '任务中通常表现更好，但样本效率相对较低。通过在统一变道决策任务上系统对比两种算法，可以分析不同'
        '算法范式在该复杂工程问题上的适用性差异。'
    )

    # ================================================================
    # 4 数据集 / 仿真环境介绍
    # ================================================================
    add_heading_styled(doc, '4 数据集 / 仿真环境介绍', level=1)

    # 4.1
    add_heading_styled(doc, '4.1 仿真平台', level=2)

    add_body_text(doc,
        '本研究不使用传统的静态标注数据集。由于强化学习的本质是通过智能体与环境交互在线学习，实验数据'
        '由 highway-env 仿真环境在训练过程中动态生成。highway-env 是一个基于简化运动学模型的高速公路'
        '多车交互仿真平台 [3]，使用 Pygame 或 EGL 进行可视化渲染，已被广泛应用于自动驾驶强化学习研究。'
    )
    add_body_text(doc,
        '环境配置为 4 车道高速公路，初始随机放置 30 辆车，每回合（episode）持续 40 秒，仿真频率 15 Hz，'
        '控制频率 5 Hz（即每 3 个仿真步执行一次决策）。车辆初始间距随机化（initial_spacing = 2 秒），'
        '车辆驶出道路或发生碰撞视为回合终止。'
    )

    # 4.2
    add_heading_styled(doc, '4.2 状态空间与观测', level=2)

    add_body_text(doc,
        '实验采用 OccupancyGrid 状态观测。将车辆周围的交通场景离散化为 16×16 的占据网格，每个网格单元'
        '编码该位置是否存在车辆，同时附加自车速度等标量特征。当系统检测到 GrayscaleImage 图像观测不可用'
        '时（如缺少 EGL 渲染支持），程序自动回退到 OccupancyGrid 观测，并切换为 MlpPolicy（多层感知机'
        '策略网络）。本实验实际运行时因环境兼容问题，所有训练均使用 OccupancyGrid + MlpPolicy，不涉及 '
        'CNN 图像训练。'
    )

    # 4.3
    add_heading_styled(doc, '4.3 动作空间', level=2)

    add_body_text(doc, '动作空间采用 DiscreteMetaAction，包含 5 种高层元动作：')

    # Table: Action space
    action_headers = ['动作编号', '名称', '含义']
    action_rows = [
        ['0', 'LANE_LEFT', '向左变道'],
        ['1', 'IDLE', '保持当前车道和速度'],
        ['2', 'LANE_RIGHT', '向右变道'],
        ['3', 'FASTER', '在当前车道加速'],
        ['4', 'SLOWER', '在当前车道减速'],
    ]
    create_styled_table(doc, action_headers, action_rows, col_widths=[2.5, 4, 6.5])

    add_paragraph_with_font(
        doc, '表1 动作空间定义',
        font_name_cn='宋体', font_name_en='Times New Roman',
        size=Pt(9), bold=False,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=Pt(3), space_after=Pt(6),
        line_spacing=1.2
    )

    # 4.4
    add_heading_styled(doc, '4.4 评价指标', level=2)

    add_body_text(doc, '每轮评估使用以下五项指标：')

    metrics = [
        ('success_rate', '成功率（未发生碰撞且正常行驶至回合结束的比例），越高越好。'),
        ('collision_rate', '碰撞率（发生碰撞的比例），越低越好。'),
        ('mean_reward', '平均 episode 累积奖励，越高越好但不能单独作为评价依据。'),
        ('mean_speed', '平均速度（m/s），接近目标速度 25 m/s 为佳。'),
        ('mean_lc', '平均每 episode 变道次数，适中为佳（过少可能不够灵活，过多意味着频繁变道）。'),
    ]
    for name, desc in metrics:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run_name = p.add_run(f'{name}：')
        set_run_font(run_name, '宋体', 'Times New Roman', Pt(12))
        run_name.bold = True
        run_desc = p.add_run(desc)
        set_run_font(run_desc, '宋体', 'Times New Roman', Pt(12))

    # 4.5
    add_heading_styled(doc, '4.5 分布外（OOD）评估场景', level=2)

    add_body_text(doc, '为评估模型在未见场景下的泛化能力，设计四种评估场景：')

    ood_scenes = [
        ('1. in_dist：', '常规分布内场景（训练配置）。'),
        ('2. lane_closure：', '车道关闭场景，在道路前方 180–210 m 区间封闭左侧两车道，迫使车辆变道避让。'),
        ('3. sudden_brake：', '前车急刹场景，在 20 秒时刻对指定前车施加 −8 m/s² 的急减速并持续 3 秒。'),
        ('4. high_density：', '高密度交通场景，将车辆数量从默认 30 辆增加到 50 辆，初始间距缩短至 1 秒。'),
    ]
    for title, desc in ood_scenes:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run_title = p.add_run(title)
        set_run_font(run_title, '宋体', 'Times New Roman', Pt(12))
        run_title.bold = True
        run_desc = p.add_run(desc)
        set_run_font(run_desc, '宋体', 'Times New Roman', Pt(12))

    add_body_text(doc, '每种场景评估 50 个 episode。')

    # ================================================================
    # 5 程序设计
    # ================================================================
    add_heading_styled(doc, '5 程序设计', level=1)

    # 5.1
    add_heading_styled(doc, '5.1 项目模块结构', level=2)

    add_body_text(doc,
        '本研究采用 Python 语言开发，基于 Stable-Baselines3 强化学习库和 highway-env 仿真平台。'
        '整体遵循模块化设计原则，主要模块如下：'
    )

    modules = [
        ('config.yaml：', '集中管理所有超参数、奖励权重、OOD 场景参数和环境配置，作为唯一配置源。'),
        ('make_env.py：', '环境工厂模块，统一环境创建流程，支持 OccupancyGrid/GrayscaleImage 自动回退、自定义奖励包装、OOD 场景包装和 Monitor 监控。'),
        ('utils/reward.py：', '自定义奖励函数模块，实现五分量加权奖励计算，包含碰撞惩罚、速度激励、变道代价、急变速代价和靠右行驶奖励，并提供 reward hacking 检测功能。'),
        ('utils/logging.py：', '日志与回调模块，包含崩溃安全的 CSVLogger、训练过程 EvalCallback（每 5000 步评估）、奖励分量记录和最优模型 checkpoint 保存。'),
        ('utils/ood_scenarios.py：', 'OOD 场景定义模块，实现车道封闭、前车急刹、高密度车流三种 Gym Wrapper。'),
        ('train_dqn.py / train_ppo.py：', '训练入口脚本，支持命令行参数配置算法、reward 类型和随机种子。'),
        ('eval.py：', '独立评估脚本，支持 OOD 场景测试和 reward hacking 检测。'),
        ('select_best.py：', '最优模型选择脚本，从每个 (algo, reward) 组合的 5 个 seed 中选择训练评估均值最高者。'),
        ('scripts/plot_curves.py / plot_seeds.py / plot_reward.py：', '可视化脚本，分别生成训练曲线图、seed 方差箱线图和奖励分量柱状图。'),
    ]
    for title, desc in modules:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run_title = p.add_run(title)
        set_run_font(run_title, '宋体', 'Times New Roman', Pt(12))
        run_title.bold = True
        run_desc = p.add_run(desc)
        set_run_font(run_desc, '宋体', 'Times New Roman', Pt(12))

    # 5.2
    add_heading_styled(doc, '5.2 奖励函数设计', level=2)

    add_body_text(doc, '奖励函数是本研究核心设计之一。整体奖励 R 为五个分量的加权和：')

    # Reward formula
    add_paragraph_with_font(
        doc, 'R = w_coll × R_coll + w_speed × R_speed + w_lc × R_lc + w_accel × R_accel + w_right × R_right',
        font_name_cn='宋体', font_name_en='Times New Roman',
        size=Pt(12), bold=False,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=Pt(3), space_after=Pt(3),
        line_spacing=1.5
    )

    reward_comps = [
        ('碰撞惩罚 R_coll：', '发生碰撞时给予 −|w_coll| 惩罚，否则为 0。这是最重要的安全约束。'),
        ('速度奖励 R_speed：', '采用高斯型函数 exp(−(v − v_target)² / (2σ²))，鼓励车速接近目标值 25 m/s。'),
        ('变道代价 R_lc：', '执行变道动作时给予 −|w_lc| 惩罚，抑制不必要的频繁变道。'),
        ('急变速代价 R_accel：', '执行加速或减速动作时给予 −|w_accel| 惩罚，提高乘坐舒适性。'),
        ('靠右奖励 R_right：', '车辆位于右侧车道时给予正向激励，符合靠右行驶的交通规则。'),
    ]
    for title, desc in reward_comps:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run_title = p.add_run(title)
        set_run_font(run_title, '宋体', 'Times New Roman', Pt(12))
        run_title.bold = True
        run_desc = p.add_run(desc)
        set_run_font(run_desc, '宋体', 'Times New Roman', Pt(12))

    add_body_text(doc, '通过调节五个权重，定义了五种驾驶风格：')

    # Table: Reward configurations
    reward_headers = ['配置', '碰撞惩罚', '速度权重', '变道代价', '急变速代价', '靠右奖励', '风格特征']
    reward_rows = [
        ['baseline', '高', '中', '低', '低', '低', '基础安全+速度'],
        ['comfort', '高', '中', '高', '高', '中', '注重舒适性'],
        ['aggressive', '中', '高', '低', '低', '低', '激进效率优先'],
        ['balanced', '高', '中', '中', '中', '中', '均衡各目标'],
        ['conservative', '极高', '低', '高', '高', '高', '极度保守'],
    ]
    create_styled_table(doc, reward_headers, reward_rows,
                        col_widths=[2.0, 1.5, 1.5, 1.5, 1.5, 1.5, 2.5])

    add_paragraph_with_font(
        doc, '表2 五种奖励函数配置及驾驶风格',
        font_name_cn='宋体', font_name_en='Times New Roman',
        size=Pt(9), bold=False,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=Pt(3), space_after=Pt(6),
        line_spacing=1.2
    )

    # 5.3
    add_heading_styled(doc, '5.3 训练与评估流程', level=2)

    add_body_text(doc, '完整的实验流程如下：')

    flow_steps = [
        ('1. 训练阶段：', '2 种算法 × 5 种 reward × 5 个 seed = 50 组训练，每组训练 50000 步，每 5000 步评估一次（10 episode），训练结束后保存 final_metrics.json。'),
        ('2. 选优阶段：', '对每个 (algo, reward) 组合，从 5 个 seed 中选出训练评估均值最高者，生成 best_selection.json，共 10 个。'),
        ('3. OOD 评估阶段：', '每个最优模型在 4 个场景下各评估 50 episode，生成 eval_all.json（10 个）和 final_result_summary.csv（40 行）。'),
        ('4. 可视化阶段：', '生成训练曲线（5 张）、seed 方差箱线图（5 张）和奖励分量柱状图（2 张），共 12 张图表。'),
    ]
    for title, desc in flow_steps:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run_title = p.add_run(title)
        set_run_font(run_title, '宋体', 'Times New Roman', Pt(12))
        run_title.bold = True
        run_desc = p.add_run(desc)
        set_run_font(run_desc, '宋体', 'Times New Roman', Pt(12))

    # 5.4
    add_heading_styled(doc, '5.4 结果追踪体系', level=2)

    add_body_text(doc, '本研究的结果通过四级 JSON/CSV 文件体系进行追踪：')

    tracking_files = [
        ('final_metrics.json：', '每组训练的最终指标（50 个文件）。'),
        ('best_selection.json：', '每组合最优模型的选择记录（10 个文件）。'),
        ('eval_all.json：', '每个最优模型的 OOD 评估完整记录（10 个文件）。'),
        ('final_result_summary.csv：', '汇总所有评估结果（40 行 × 12 列）。'),
    ]
    for title, desc in tracking_files:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run_title = p.add_run(title)
        set_run_font(run_title, '宋体', 'Times New Roman', Pt(12))
        run_title.bold = True
        run_desc = p.add_run(desc)
        set_run_font(run_desc, '宋体', 'Times New Roman', Pt(12))

    add_body_text(doc, '这一文件体系确保了从原始训练到最终结论的完整可追溯性。')

    # ================================================================
    # 6 实验结果与分析
    # ================================================================
    add_heading_styled(doc, '6 实验结果与分析', level=1)

    # 6.1
    add_heading_styled(doc, '6.1 实验设置', level=2)

    add_body_text(doc,
        '所有实验在纯 CPU 环境下运行。每组训练 50000 步，每 5000 步进行一次中间评估（10 episode）。'
        '每个 (algo, reward) 组合使用 5 个随机种子（42, 123, 456, 789, 1024）进行重复实验以保证统计'
        '可靠性。最终从 10 个最优模型中选取代表性的结果进行分析。'
    )

    # 6.2
    add_heading_styled(doc, '6.2 训练曲线分析', level=2)

    add_body_text(doc, '训练曲线反映了 DQN 与 PPO 在不同 reward 配置下的学习过程。')

    # Figures 1-3: Training curves
    add_figure_with_caption(doc,
        os.path.join(FIGURES_DIR, 'training_curves_aggressive.png'),
        '图1 aggressive 配置下 DQN 与 PPO 训练曲线对比')
    add_figure_with_caption(doc,
        os.path.join(FIGURES_DIR, 'training_curves_baseline.png'),
        '图2 baseline 配置下 DQN 与 PPO 训练曲线对比')
    add_figure_with_caption(doc,
        os.path.join(FIGURES_DIR, 'training_curves_balanced.png'),
        '图3 balanced 配置下 DQN 与 PPO 训练曲线对比')

    add_body_text(doc,
        '从曲线中可以观察到：PPO 算法的收敛速度明显快于 DQN。在 aggressive、baseline 和 balanced '
        '三种配置下，PPO 在约 15000–20000 步后即达到较高且稳定的平均 episode reward，而 DQN 在 50000 '
        '步内仍呈现缓慢上升或剧烈震荡趋势。这一差异与算法特性有关：PPO 每次使用最新策略采集的数据进行'
        '多轮优化更新，能更快地适应环境；DQN 作为 off-policy 算法，需要积累足够的经验回放数据后才能有效'
        '学习，且 OccupancyGrid 观测空间较大增加了 Q 函数估计的难度。'
    )

    # 6.3
    add_heading_styled(doc, '6.3 Seed 方差分析', level=2)

    add_figure_with_caption(doc,
        os.path.join(FIGURES_DIR, 'seed_variance_aggressive.png'),
        '图4 aggressive 配置下 DQN 与 PPO 的 Seed 方差箱线图')

    add_body_text(doc,
        'Seed 方差箱线图展示了各 (algo, reward) 组合在 5 个不同随机种子下的性能分布。PPO 的箱体宽度'
        '（四分位距）普遍小于 DQN，中位数也更高，说明 PPO 对随机种子的敏感性较低，训练过程更加稳定。'
        '这一优势来源于 PPO 的 clipped surrogate objective，该机制天然限制了策略更新幅度，减少了训练的'
        '随机波动。'
    )

    # 6.4
    add_heading_styled(doc, '6.4 奖励分量分析', level=2)

    add_figure_with_caption(doc,
        os.path.join(FIGURES_DIR, 'reward_components_ppo.png'),
        '图5 PPO 在不同 reward 配置下的奖励分量对比')
    add_figure_with_caption(doc,
        os.path.join(FIGURES_DIR, 'reward_components_dqn.png'),
        '图6 DQN 在不同 reward 配置下的奖励分量对比')

    add_body_text(doc,
        '奖励分量柱状图揭示了两种算法在不同 reward 配置下各奖励分量的贡献情况。PPO 的碰撞惩罚分量始终'
        '保持在较低水平（意味着碰撞较少），而 DQN 的碰撞惩罚分量很高。PPO aggressive 的速度奖励分量显著'
        '高于其他配置，这与 aggressive 配置中较高的速度权重一致。'
    )

    # 6.5
    add_heading_styled(doc, '6.5 最终平均结果', level=2)

    add_body_text(doc, '以下为四个场景下的平均结果汇总表（按成功率降序排列）：')

    # Table 3: Final average results
    result_headers = ['排名', '算法', '奖励配置', '成功率', '碰撞率', '平均回报', '平均速度', '变道次数']
    result_rows = [
        ['1', 'PPO', 'aggressive', '0.975', '0.025', '163.40', '20.36', '14.10'],
        ['2', 'PPO', 'baseline', '0.885', '0.115', '75.19', '20.73', '16.12'],
        ['3', 'PPO', 'balanced', '0.795', '0.205', '66.55', '21.67', '0.10'],
        ['4', 'PPO', 'comfort', '0.795', '0.205', '66.55', '21.67', '0.10'],
        ['5', 'DQN', 'comfort', '0.085', '0.915', '40.44', '24.98', '0.00'],
        ['6', 'DQN', 'conservative', '0.085', '0.915', '1.53', '24.98', '0.00'],
        ['7', 'PPO', 'conservative', '0.085', '0.915', '1.53', '24.98', '0.00'],
        ['8', 'DQN', 'baseline', '0.060', '0.940', '24.72', '25.38', '68.07'],
        ['9', 'DQN', 'aggressive', '0.020', '0.980', '64.21', '24.96', '66.59'],
        ['10', 'DQN', 'balanced', '0.000', '1.000', '4.79', '29.54', '0.00'],
    ]

    table3 = doc.add_table(rows=1 + len(result_rows), cols=len(result_headers))
    table3.style = 'Table Grid'
    table3.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Find which values to bold (best in each column)
    bold_rows_cols = {
        (0, 3), (0, 4), (0, 5),  # PPO aggressive: success, collision, reward
        (7, 7),  # DQN baseline: mean_lc
        (8, 7),  # DQN aggressive: mean_lc
    }

    # Header
    for i, header in enumerate(result_headers):
        cell = table3.rows[0].cells[i]
        set_cell_font(cell, header, size=Pt(8), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3"/>')
        cell._element.get_or_add_tcPr().append(shading)

    # Body
    for r, row_data in enumerate(result_rows):
        for c, cell_text in enumerate(row_data):
            cell = table3.rows[r + 1].cells[c]
            is_bold = (r, c) in bold_rows_cols
            set_cell_font(cell, cell_text, size=Pt(8), bold=is_bold,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Set column widths for result table
    result_col_widths = [1.0, 1.2, 2.0, 1.5, 1.5, 1.5, 1.5, 1.5]
    for row in table3.rows:
        for i, width in enumerate(result_col_widths):
            row.cells[i].width = Cm(width)

    add_paragraph_with_font(
        doc, '表3 四个场景下各模型平均结果（按成功率降序排列，加粗为各列最优值）',
        font_name_cn='宋体', font_name_en='Times New Roman',
        size=Pt(9), bold=False,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=Pt(3), space_after=Pt(3),
        line_spacing=1.2
    )

    add_body_text(doc,
        '需要说明的是：PPO balanced 与 PPO comfort 的最优种子均为 seed 456，且在该种子下两个模型的表现'
        '完全一致（均由高变道代价导致变道次数接近于零），因此二者在汇总表中数据相同。PPO conservative 与 '
        'DQN conservative 的最优种子均为 seed 789，在该种子下两模型同样表现出高度一致的指标。'
    )

    # 6.6
    add_heading_styled(doc, '6.6 重点分析', level=2)

    # 6.6 sub-analyses as Heading 3
    add_heading_styled(doc, 'PPO aggressive —— 综合最佳', level=3)

    add_body_text(doc,
        'PPO aggressive 在四个场景下的平均成功率达到 97.5%，碰撞率仅 2.5%，平均回报 163.40，是综合性能'
        '最佳的策略。aggressive 配置中碰撞惩罚权重适中、速度权重较高，使得 PPO 在学习过程中积极探索变道'
        '行为以提高速度，同时 PPO 的 clipped objective 机制确保策略更新不过度激进，在探索与安全之间取得了'
        '良好平衡。PPO aggressive 在三种 OOD 场景下均保持了稳定表现：高密度场景成功率 98%、急刹场景 96%、'
        '车道关闭场景 98%，展现了优秀的泛化能力。'
    )

    add_heading_styled(doc, 'PPO balanced / comfort —— 保守驾驶', level=3)

    add_body_text(doc,
        'PPO balanced 和 PPO comfort 表现出几乎一致的行为特征：平均每 episode 变道仅 0.10 次，几乎不主动'
        '变道。这两种配置对变道动作和急变速动作施加了更强的惩罚，导致模型倾向于保持当前车道行驶。优点是'
        '乘坐舒适性高（变道和急变速极少），缺点是在真正需要变道避让时表现不足，碰撞率（20.5%）明显高于 '
        'aggressive（2.5%），且平均速度（21.67 m/s）略低于目标速度。这说明过度惩罚变道行为会抑制模型在'
        '必要时采取变道操作的能力。'
    )

    add_heading_styled(doc, 'DQN baseline / aggressive —— 频繁变道', level=3)

    add_body_text(doc,
        'DQN baseline 的平均变道次数高达 68.07 次/episode，DQN aggressive 为 66.59 次/episode，而 PPO '
        'aggressive 仅为 14.10 次/episode。这种异常的频繁变道伴随极高的碰撞率（94%–98%），说明 DQN 并非'
        '学到了安全的变道策略，而是在利用 reward 函数中的某些漏洞（reward hacking）——通过高频变道获取'
        '速度奖励，却无法避免因此导致的碰撞。这是 DQN 在本实验设定下策略不稳定的典型表现。'
    )

    add_heading_styled(doc, 'DQN balanced —— 完全失败', level=3)

    add_body_text(doc,
        'DQN balanced 的碰撞率达到 100%，即每个 episode 都发生碰撞。同时其平均速度高达 29.54 m/s'
        '（远超目标速度 25 m/s），说明模型只学会了加速，完全未学到有效的变道或减速避让行为。这一结果'
        '表明，当 reward 函数中各分量权重设置对 DQN 不利时，DQN 可能完全无法学到有效策略。'
    )

    add_heading_styled(doc, 'PPO / DQN conservative —— 过于保守导致失败', level=3)

    add_body_text(doc,
        'PPO conservative 和 DQN conservative 的碰撞率均为 91.5%，成功率仅 8.5%。conservative 配置中'
        '碰撞惩罚极高但速度权重很低，模型在过于严苛的惩罚下未能有效探索，反而导致了更差的安全表现。'
        '这验证了一个重要结论：过于保守的 reward 设计可能适得其反。'
    )

    # 6.7
    add_heading_styled(doc, '6.7 核心发现', level=2)

    add_body_text(doc,
        '高回报不等于安全策略。DQN aggressive 的平均回报为 64.21，高于 PPO conservative 的 1.53，但 DQN '
        'aggressive 的碰撞率高达 98%，是极度危险的策略。评价模型时必须结合 success_rate（成功率）、'
        'collision_rate（碰撞率）、mean_lc（变道次数）和 mean_speed（速度）进行多维度综合评价，不能只看 '
        'mean_reward。'
    )

    add_body_text(doc,
        'PPO 在本任务中显著优于 DQN。PPO 的全部五种 reward 配置在平均排名上均位于 DQN 对应配置之上'
        '（唯一的例外是 conservative，两者表现同样差）。PPO 的成功得益于其 clipped objective 提供的训练'
        '稳定性，以及 on-policy 更新机制在 OccupancyGrid 观测下的更好适应能力。'
    )

    add_body_text(doc,
        '多算法、多 reward、多场景实验设计的必要性。本实验充分验证了系统对比多种方法、多种配置和多种评估'
        '场景的重要性。如果仅测试一种算法或一种 reward 配置，可能得出片面的结论。'
    )

    # ================================================================
    # 7 工程问题复杂性与社会文化因素分析
    # ================================================================
    add_heading_styled(doc, '7 工程问题复杂性与社会文化因素分析', level=1)

    add_body_text(doc,
        '本部分对应课程评分中"工程与社会"维度（占比 10%），分析自动驾驶变道决策涉及的工程复杂性、'
        '社会文化因素和跨文化国际交流能力。'
    )

    # 7.1
    add_heading_styled(doc, '7.1 安全责任归属问题', level=2)

    add_body_text(doc,
        '自动驾驶变道决策系统的安全责任归属是一个复杂的法律和伦理问题。当自动驾驶系统做出变道决策导致'
        '交通事故时，责任可能涉及多方主体：车辆驾驶员（如未能及时接管）、汽车制造商（如系统设计缺陷）、'
        '算法开发者（如模型缺陷导致错误决策）、传感器供应商（如感知数据错误）以及基础设施提供者（如道路'
        '标线不清）。目前各国尚未形成统一的自动驾驶法律责任框架。部分国家已开始推进自动驾驶伦理和安全'
        '监管规则建设，中国也正在积极推进自动驾驶相关立法。本研究中，即使表现最好的 PPO aggressive 策略'
        '碰撞率仍为 2.5%，这意味着在仿真环境中每 40 次变道决策就可能发生一次碰撞——这在真实交通中是不可'
        '接受的。真实道路部署对安全性的要求远高于仿真实验中的成功率水平，模型从 97.5% 成功率到实际部署'
        '要求之间的差距，体现了从学术研究到工程应用的巨大鸿沟。'
    )

    # 7.2
    add_heading_styled(doc, '7.2 不同国家和地区的交通法规差异', level=2)

    add_body_text(doc,
        '全球各国的交通法规存在显著差异，直接影响自动驾驶变道决策系统的设计和部署。以下为几个典型差异：'
    )

    legal_diffs = [
        ('通行方向：', '中国、美国等多数国家采用右侧通行，而英国、日本、澳大利亚、印度等约 75 个国家和地区采用左侧通行。本研究的靠右奖励（right_lane reward）在左侧通行国家必须修改为靠左奖励。'),
        ('限速规定：', '德国部分高速公路（Autobahn）无限速或仅有建议速度，中国高速公路普遍限速 120 km/h，美国各州限速 65–85 mph 不等。模型的速度目标和安全距离参数需要根据部署地区调整。'),
        ('变道规则：', '不同国家对变道信号灯使用时长、变道最小安全间距、实线禁止变道等规则存在差异。模型需要遵守当地交通法规。'),
    ]
    for title, desc in legal_diffs:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run_title = p.add_run(title)
        set_run_font(run_title, '宋体', 'Times New Roman', Pt(12))
        run_title.bold = True
        run_desc = p.add_run(desc)
        set_run_font(run_desc, '宋体', 'Times New Roman', Pt(12))

    # 7.3
    add_heading_styled(doc, '7.3 驾驶文化差异', level=2)

    add_body_text(doc,
        '不同文化背景下，人们对激进驾驶和保守驾驶的接受程度存在显著差异。在南欧和部分拉美国家，较小的'
        '跟车距离和较频繁的变道被普遍接受；在北欧和日本，更注重平稳驾驶和较大的安全距离。中国城市交通中'
        '常见的"加塞"和频繁鸣笛行为在德国或日本会被视为不礼貌甚至违规。本研究中的 aggressive 配置在追求'
        '效率的同时保持了较低的碰撞率，但在驾驶文化保守的地区可能被认为不够舒适；comfort 和 balanced '
        '配置虽然更平稳，但在需要果断变道的场景下反应不足。自动驾驶系统需要能够适配不同地区的驾驶文化'
        '期望，这不仅是技术问题，也是跨文化理解问题。'
    )

    # 7.4
    add_heading_styled(doc, '7.4 乘客舒适性与交通效率的权衡', level=2)

    add_body_text(doc,
        '本实验中 balanced/comfort 配置（mean_lc 约等于 0）和 aggressive 配置（mean_lc 约等于 14）的对比，'
        '直接体现了舒适性与效率之间的经典权衡。在共享出行场景下，乘客可能更注重舒适性；在物流运输场景下，'
        '运营方可能更注重效率。自动驾驶系统需要能够根据不同应用场景和用户偏好调整驾驶策略，实现个性化配置。'
    )

    # 7.5
    add_heading_styled(doc, '7.5 标准规范与法规合规', level=2)

    add_body_text(doc,
        '自动驾驶系统在真实道路部署前需要满足多项国际标准和安全法规。虽然具体标准编号可能因时效性而变化，'
        '但通常需要参考功能安全标准（如 ISO 26262 道路车辆功能安全）、预期功能安全标准（如 ISO 21448 '
        'SOTIF）、网络安全标准（如 ISO/SAE 21434）以及自动驾驶测试场景相关标准。机器学习模型的可解释性'
        '也是合规的关键挑战：深度神经网络作为"黑箱"模型，其变道决策的内部推理过程难以向监管机构清晰解释。'
        '本研究目前处于仿真实验阶段，尚未涉及上述标准的正式验证。'
    )

    # 7.6
    add_heading_styled(doc, '7.6 国际视野与跨文化交流', level=2)

    add_body_text(doc,
        '自动驾驶是全球性研究课题，涉及跨国际的学术交流和技术合作。不同国家和地区的技术路线选择受到当地'
        '产业政策、道路基础设施、数据法规和市场需求的影响：例如北美企业在 L4 级 Robotaxi 领域投入较大，'
        '中国企业强调车路协同技术路线，欧洲企业则在功能安全和法规合规方面具有传统优势。本研究的英文摘要'
        '和国际化参考文献列表也是对基本跨文化国际交流能力的体现。'
    )

    # 7.7
    add_heading_styled(doc, '7.7 低成本部署与可持续发展', level=2)

    add_body_text(doc,
        '在真实道路部署自动驾驶变道系统还需考虑：计算资源约束（车载芯片算力有限，需要模型轻量化）、'
        '传感器成本（不同感知方案的成本差异可达数十倍）、维护和 OTA 升级成本、数据隐私保护（车辆轨迹数据'
        '的收集和使用需符合 GDPR、中国《个人信息保护法》等规定）。此外，自动驾驶的普及可能对交通运输就业'
        '结构产生深远影响——职业司机岗位可能减少，同时新增自动驾驶运维、远程监控等新型技术岗位。在技术'
        '推广过程中需要关注社会公平和劳动力转型问题。'
    )

    # ================================================================
    # 8 结论
    # ================================================================
    add_heading_styled(doc, '8 结论', level=1)

    add_body_text(doc,
        '本研究基于 highway-env 仿真平台，将高速公路自主变道决策建模为强化学习中的序贯决策问题，系统'
        '比较了 DQN 和 PPO 两种深度强化学习算法在五种多目标奖励函数配置下的性能，并在三种 OOD 场景下评估'
        '了模型的泛化能力。主要结论如下：'
    )

    conclusions = [
        'PPO 算法在收敛速度、训练稳定性和最终性能方面整体优于 DQN，更适合高速公路变道决策任务。PPO 的 '
        'clipped surrogate objective 机制有效限制了策略更新幅度，减少了训练波动，在 OccupancyGrid 观测下'
        '展现出更好的适应能力。',

        'PPO aggressive 策略在四个场景下的平均成功率为 97.5%、碰撞率为 2.5%、平均回报为 163.40，是本研究'
        '综合性能最优的策略，展现了良好的安全性和泛化能力。',

        'PPO balanced 和 PPO comfort 策略的变道次数极少（mean_lc = 0.10），体现了保守、舒适的驾驶风格，'
        '但其碰撞率（20.5%）高于 aggressive（2.5%），说明在真正需要变道时不够灵活。',

        'DQN 在本实验设定下存在严重的策略不稳定问题：DQN baseline 和 DQN aggressive 出现频繁变道'
        '（mean_lc > 60）和高碰撞率（> 94%）的 reward hacking 现象；DQN balanced 碰撞率达到 100%。'
        '这说明 DQN 在 OccupancyGrid + MlpPolicy 的组合下难以学到安全有效的高速公路驾驶策略。',

        '高回报不一定等于安全策略。评价模型时必须综合 success_rate、collision_rate、mean_lc 和 mean_speed '
        '等多个指标，避免仅凭 mean_reward 做出误导性结论。',

        '多算法、多 reward 配置、多 OOD 场景的系统实验设计对于全面评估自动驾驶决策模型至关重要。',
    ]
    for i, conc in enumerate(conclusions):
        add_body_text(doc, f'（{i+1}）{conc}')

    add_body_text(doc, '后续改进方向：')
    future_work = [
        '引入更真实的交通流模型（如 SUMO 联合仿真），缩小仿真与现实差距；',
        '增加多目标安全约束，将碰撞率优化目标直接纳入训练过程；',
        '引入可解释性分析方法（如注意力可视化、SHAP 值分析），提升模型的透明度和可信度；',
        '使用多智能体强化学习方法，建模多车协同决策；',
        '探索安全强化学习方法（如 Constrained RL、Safe RL），在训练过程中显式施加安全约束。',
    ]
    for i, item in enumerate(future_work):
        add_body_text(doc, f'（{i+1}）{item}')

    # ================================================================
    # 参考文献
    # ================================================================
    add_heading_styled(doc, '参考文献', level=1)

    references = [
        '[1] Mnih V, Kavukcuoglu K, Silver D, et al. Human-level control through deep reinforcement learning[J]. Nature, 2015, 518(7540): 529-533.',
        '[2] Schulman J, Wolski F, Dhariwal P, et al. Proximal policy optimization algorithms[J]. arXiv preprint arXiv:1707.06347, 2017.',
        '[3] Leurent E. An Environment for Autonomous Driving Decision-Making[EB/OL]. https://github.com/Farama-Foundation/HighwayEnv, 2018.',
        '[4] Sutton R S, Barto A G. Reinforcement Learning: An Introduction[M]. 2nd ed. MIT Press, 2018.',
        '[5] Kiran B R, Sobh I, Talpaert V, et al. Deep reinforcement learning for autonomous driving: A survey[J]. IEEE Transactions on Intelligent Transportation Systems, 2021, 23(6): 4909-4926.',
        '[6] Codevilla F, Muller M, Lopez A, et al. End-to-end driving via conditional imitation learning[C]. IEEE International Conference on Robotics and Automation (ICRA), 2018.',
        '[7] Kendall A, Hawke J, Janz D, et al. Learning to drive in a day[C]. IEEE International Conference on Robotics and Automation (ICRA), 2019.',
        '[8] 中华人民共和国道路交通安全法[S]. 2021 年修订版.',
        '[9] Ding Z, Huang H. Autonomous driving decision-making with deep reinforcement learning: A review[J]. IEEE Access, 2023, 11: 10590-10609.',
    ]
    for ref in references:
        add_paragraph_with_font(
            doc, ref,
            font_name_cn='宋体', font_name_en='Times New Roman',
            size=Pt(10.5), bold=False,
            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            first_line_indent=None,
            space_before=Pt(1), space_after=Pt(1),
            line_spacing=1.3
        )

    # ================================================================
    # Add page numbers
    # ================================================================
    add_page_number(doc)

    # ================================================================
    # Save
    # ================================================================
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Report saved to: {OUTPUT_PATH}")
    print("Done!")


if __name__ == '__main__':
    create_report()
