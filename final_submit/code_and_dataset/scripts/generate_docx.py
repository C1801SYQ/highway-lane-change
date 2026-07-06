"""Generate Word document from final_report.md for course submission."""
import re
import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PROJECT_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
REPORT_MD = os.path.join(PROJECT_DIR, "docs", "final_report.md")
OUTPUT_DOCX = os.path.join(PROJECT_DIR, "final_submit", "基于DQN与PPO的高速公路自主变道决策研究.docx")
FIGURES_DIR = os.path.join(PROJECT_DIR, "figures")

# Figure mapping: markdown reference -> actual file
FIGURE_MAP = {
    "figures/training_curves_aggressive.png": "training_curves_aggressive.png",
    "figures/training_curves_baseline.png": "training_curves_baseline.png",
    "figures/training_curves_balanced.png": "training_curves_balanced.png",
    "figures/seed_variance_aggressive.png": "seed_variance_aggressive.png",
    "figures/reward_components_ppo.png": "reward_components_ppo.png",
    "figures/reward_components_dqn.png": "reward_components_dqn.png",
}

def set_run_font(run, cn_font="宋体", en_font="Times New Roman", size=Pt(12), bold=False):
    """Set font for a run with both Chinese and English fonts."""
    run.font.size = size
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), cn_font)
    rFonts.set(qn('w:ascii'), en_font)
    rFonts.set(qn('w:hAnsi'), en_font)
    rPr.insert(0, rFonts)

def set_paragraph_spacing(paragraph, line_spacing=1.5, space_before=0, space_after=0):
    """Set paragraph spacing."""
    pf = paragraph.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)

def add_heading_styled(doc, text, level=1):
    """Add a heading with proper Chinese formatting."""
    p = doc.add_paragraph()
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)

    if level == 1:
        set_run_font(run, cn_font="黑体", size=Pt(16), bold=True)  # 三号
        set_paragraph_spacing(p, line_spacing=1.5, space_before=12, space_after=6)
    elif level == 2:
        set_run_font(run, cn_font="黑体", size=Pt(14), bold=True)  # 四号
        set_paragraph_spacing(p, line_spacing=1.5, space_before=10, space_after=4)
    elif level == 3:
        set_run_font(run, cn_font="黑体", size=Pt(12), bold=True)  # 小四加粗
        set_paragraph_spacing(p, line_spacing=1.5, space_before=8, space_after=2)

    return p

def add_body_paragraph(doc, text, bold=False, alignment=None):
    """Add a body paragraph with 小四 font."""
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    run = p.add_run(text)
    set_run_font(run, size=Pt(12), bold=bold)  # 小四
    set_paragraph_spacing(p, line_spacing=1.5, space_before=0, space_after=0)
    return p

def add_image_centered(doc, image_path, width_inches=5.0):
    """Add a centered image."""
    full_path = os.path.join(FIGURES_DIR, image_path)
    if os.path.exists(full_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(full_path, width=Inches(width_inches))
        set_paragraph_spacing(p, line_spacing=1.5, space_before=6, space_after=6)
        # Add caption
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = caption.add_run(f"图：{image_path.replace('.png', '').replace('_', ' ').title()}")
        set_run_font(cap_run, size=Pt(10), bold=False)
        set_paragraph_spacing(caption, line_spacing=1.5, space_before=0, space_after=6)

def add_table_from_md(doc, lines):
    """Parse and add a markdown table."""
    if len(lines) < 2:
        return

    # Parse header
    headers = [h.strip() for h in lines[0].split('|')[1:-1]]
    # Skip separator line
    rows_data = []
    for line in lines[2:]:
        if line.startswith('|'):
            cells = [c.strip() for c in line.split('|')[1:-1]]
            rows_data.append(cells)

    if not headers:
        return

    table = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        set_run_font(run, size=Pt(10), bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for row_idx, row_data in enumerate(rows_data):
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < len(headers):
                cell = table.rows[row_idx + 1].cells[col_idx]
                cell.text = ''
                p = cell.paragraphs[0]
                run = p.add_run(cell_text)
                set_run_font(run, size=Pt(10), bold=False)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add spacing after table
    p = doc.add_paragraph()
    set_paragraph_spacing(p, line_spacing=1.5, space_before=3, space_after=3)

def generate_docx():
    """Main function to generate the Word document."""
    os.makedirs(os.path.dirname(OUTPUT_DOCX), exist_ok=True)

    doc = Document()

    # Set default margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    # Read the markdown file
    with open(REPORT_MD, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')

    i = 0
    in_table = False
    table_lines = []
    in_code_block = False

    while i < len(lines):
        line = lines[i]

        # Skip horizontal rules and frontmatter markers
        if line.strip() == '---':
            i += 1
            continue

        # Handle code blocks
        if line.strip().startswith('```'):
            in_code_block = not in_code_block
            i += 1
            continue

        if in_code_block:
            i += 1
            continue

        # Handle headings
        if line.startswith('# ') and not line.startswith('## '):
            text = line[2:].strip()
            add_heading_styled(doc, text, level=1)
            i += 1
            continue

        if line.startswith('## '):
            text = line[3:].strip()
            add_heading_styled(doc, text, level=2)
            i += 1
            continue

        if line.startswith('### '):
            text = line[4:].strip()
            add_heading_styled(doc, text, level=3)
            i += 1
            continue

        # Handle image references
        img_match = re.match(r'【插图：(.+?)】', line.strip())
        if img_match:
            img_ref = img_match.group(1).strip()
            if img_ref in FIGURE_MAP:
                add_image_centered(doc, FIGURE_MAP[img_ref], width_inches=5.0)
            i += 1
            continue

        # Handle tables
        if line.startswith('|') and line.rstrip().endswith('|'):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        else:
            if in_table:
                add_table_from_md(doc, table_lines)
                table_lines = []
                in_table = False
                # Don't increment i, process current line

        # Handle bold text (**text**)
        if '**' in line:
            # Split by bold markers and create runs
            p = doc.add_paragraph()
            set_paragraph_spacing(p, line_spacing=1.5, space_before=0, space_after=0)
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    set_run_font(run, size=Pt(12), bold=True)
                else:
                    run = p.add_run(part)
                    set_run_font(run, size=Pt(12), bold=False)
            i += 1
            continue

        # Handle empty lines
        if not line.strip():
            i += 1
            continue

        # Regular paragraph
        # Clean up markdown formatting
        clean_line = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
        clean_line = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', clean_line)
        clean_line = re.sub(r'`(.+?)`', r'\1', clean_line)

        # Skip figure reference lines and other non-content
        if clean_line.strip().startswith('>') or clean_line.strip().startswith('*分析报告结束*'):
            i += 1
            continue

        # Detect bullet points
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            p = doc.add_paragraph()
            run = p.add_run('• ' + text)
            set_run_font(run, size=Pt(12), bold=False)
            set_paragraph_spacing(p, line_spacing=1.5, space_before=0, space_after=0)
            p.paragraph_format.left_indent = Cm(0.75)
            i += 1
            continue

        # Add as body paragraph
        if clean_line.strip():
            p = doc.add_paragraph()
            run = p.add_run(clean_line.strip())
            set_run_font(run, size=Pt(12), bold=False)
            set_paragraph_spacing(p, line_spacing=1.5, space_before=0, space_after=0)

        i += 1

    # Handle any remaining table
    if in_table and table_lines:
        add_table_from_md(doc, table_lines)

    # Save
    doc.save(OUTPUT_DOCX)
    print(f"Word document saved to: {OUTPUT_DOCX}")

if __name__ == "__main__":
    generate_docx()
